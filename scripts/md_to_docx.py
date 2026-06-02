"""
Convert DOCUMENTATION.md to DOCUMENTATION.docx.

Handles the Markdown subset used in the doc:
  * ATX headings (#, ##, ###, ####)
  * horizontal rules (---)
  * pipe tables (with a header separator row)
  * fenced code blocks (``` ... ```), including language hints
  * blockquotes (>), with nested code/lists rendered as plain lines
  * unordered (- ) and ordered (1. ) lists, including nested indentation
  * inline formatting: **bold**, `code`
Run:
    .venv/Scripts/python.exe scripts/md_to_docx.py
"""

from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor, Inches

SRC = "DOCUMENTATION.md"
OUT = "DOCUMENTATION.docx"

MONO = "Consolas"
CODE_BG_SHADE = "F2F2F2"
CODE_GREY = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0x2F, 0x54, 0x96)


# ---------------------------------------------------------------------------
# Inline formatting: **bold** and `code`
# ---------------------------------------------------------------------------
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str):
    """Add runs to a paragraph, honoring **bold** and `inline code`."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO
            run.font.color.rgb = CODE_GREY
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Code block shading helper
# ---------------------------------------------------------------------------
def _shade(paragraph, fill: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _shade(p, CODE_BG_SHADE)
        run = p.add_run(line if line else " ")
        run.font.name = MONO
        run.font.size = Pt(9)
        run.font.color.rgb = CODE_GREY


# ---------------------------------------------------------------------------
# Table helper
# ---------------------------------------------------------------------------
def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc: Document, rows: list[list[str]]):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].text = ""
        run = cell.paragraphs[0].add_run(text)
        run.bold = True

    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if i >= len(cells):
                break
            cells[i].paragraphs[0].text = ""
            add_inline_runs(cells[i].paragraphs[0], text)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------
def convert(src: str, out: str):
    with open(src, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            add_code_block(doc, code)
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            heading.text = ""
            add_inline_runs(heading, m.group(2))
            i += 1
            continue

        # Pipe table
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\|[\s:\-|]+\|$", lines[i + 1].strip()):
            rows = [_split_row(stripped)]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_row(lines[i].strip()))
                i += 1
            add_table(doc, rows)
            continue

        # Blockquote (may span multiple lines)
        if stripped.startswith(">"):
            quote_lines = []
            while i < n and lines[i].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            text = " ".join(l.strip() for l in quote_lines if l.strip())
            p = doc.add_paragraph(style="Intense Quote")
            add_inline_runs(p, text)
            continue

        # Unordered list item (supports one level of nesting via indent)
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent // 2))
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Ordered list item
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style="List Number")
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.5 + 0.25 * (indent // 2))
            add_inline_runs(p, m.group(2))
            i += 1
            continue

        # Blank line
        if stripped == "":
            i += 1
            continue

        # Plain paragraph (collect consecutive non-special lines)
        para_lines = [stripped]
        i += 1
        while i < n:
            nxt = lines[i].strip()
            if (nxt == "" or nxt.startswith(("#", ">", "|", "```", "---"))
                    or re.match(r"^(\s*)[-*]\s+", lines[i])
                    or re.match(r"^(\s*)\d+\.\s+", lines[i])):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))

    doc.save(out)
    print(f"Wrote {out} ({os.path.getsize(out):,} bytes)")


if __name__ == "__main__":
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    convert(os.path.join(base, SRC), os.path.join(base, OUT))
